"""
OCR and text extraction for GraveKeeper
"""
import pytesseract
from PIL import Image
import PyPDF2
import docx
import pandas as pd
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import io
import fitz  # PyMuPDF for better PDF handling
import openpyxl
from openpyxl import load_workbook

from src.config.settings import settings
from src.utils.logger import logger
from src.utils.file_utils import get_file_extension, get_file_category

class TextExtractor:
    """Extract text from various file types using OCR and native extraction"""

    def __init__(self, pdf_ocr: bool = True):
        """Initialize text extractor
        Args:
            pdf_ocr: Whether to run OCR on PDF images (default True)
        """
        self.pdf_ocr = pdf_ocr
        # Configure tesseract path if needed
        # pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
        
    def extract_text(self, file_path: Path) -> Dict:
        """
        Extract text from file based on its type
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            file_extension = get_file_extension(str(file_path))
            file_category = get_file_category(str(file_path))
            
            result = {
                'file_path': str(file_path),
                'file_extension': file_extension,
                'file_category': file_category,
                'text_content': '',
                'extraction_method': '',
                'confidence': 0.0,
                'error': None,
                'metadata': {}
            }
            
            if file_category == 'images':
                result.update(self._extract_from_image(file_path))
            elif file_category == 'documents':
                result.update(self._extract_from_document(file_path))
            elif file_category == 'spreadsheets':
                result.update(self._extract_from_spreadsheet(file_path))
            elif file_category == 'presentations':
                result.update(self._extract_from_presentation(file_path))
            else:
                result['error'] = f'Unsupported file type: {file_extension}'
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {
                'file_path': str(file_path),
                'error': str(e),
                'text_content': '',
                'extraction_method': 'failed'
            }
    
    def _extract_from_image(self, file_path: Path) -> Dict:
        """
        Extract text from image using OCR
        
        Args:
            file_path: Path to image file
            
        Returns:
            Dictionary with extracted text and OCR metadata
        """
        try:
            # Open image
            image = Image.open(file_path)
            # Use Qwen2.5VL vision model for description
            from src.ai.providers.qwen2_5vl import Qwen2_5VLProvider
            qwen_provider = Qwen2_5VLProvider()
            vision_result = qwen_provider.analyze_vision(image)
            vision_text = vision_result.get('explanation', str(vision_result))
            return {
                'text_content': vision_text.strip(),
                'extraction_method': 'qwen2.5vl-vision',
                'confidence': 1.0 if vision_text.strip() else 0.0,
                'metadata': {
                    'image_size': image.size,
                    'image_mode': image.mode,
                    'qwen2.5vl_vision_result': vision_result
                }
            }
        except Exception as e:
            logger.error(f"Qwen2.5VL vision extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'qwen2.5vl-vision_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_document(self, file_path: Path) -> Dict:
        """
        Extract text from document files (PDF, DOC, DOCX, TXT)
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary with extracted text
        """
        file_extension = get_file_extension(str(file_path))
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self._extract_from_word(file_path)
        elif file_extension == '.txt':
            return self._extract_from_text(file_path)
        else:
            return {
                'text_content': '',
                'extraction_method': 'unsupported',
                'error': f'Unsupported document type: {file_extension}'
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            text_content = ""
            ocr_text = ""
            extraction_method = ""
            pdf_pages = 'unknown'
            doc = None
            pymupdf_success = False
            try:
                doc = fitz.open(file_path)
                pdf_pages = len(doc)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                extraction_method = "pymupdf"
                pymupdf_success = True
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {file_path}, trying PyPDF2: {e}")
                # Fallback to PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() or ""
                    extraction_method = "pypdf2"
                    pdf_pages = len(pdf_reader.pages) if 'pdf_reader' in locals() else 'unknown'
                except Exception as e2:
                    logger.warning(f"PyPDF2 also failed for {file_path}: {e2}")
            # Always run OCR on PDF images unless disabled
            qwen2_5vl_results = []
            if self.pdf_ocr:
                try:
                    from src.ai.providers.qwen2_5vl import Qwen2_5VLProvider
                    qwen_provider = Qwen2_5VLProvider()
                except Exception as e:
                    logger.error(f"Qwen2.5VL import failed: {e}")
                    qwen_provider = None
                try:
                    if doc is None:
                        doc = fitz.open(file_path)
                        pdf_pages = len(doc)
                    vision_descriptions = []
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap()
                        img = Image.open(io.BytesIO(pix.tobytes()))
                        ocr_result = pytesseract.image_to_string(
                            img,
                            lang=settings.OCR_LANGUAGE,
                            config='--psm 6'
                        )
                        ocr_text += ocr_result + "\n"
                        # Run Qwen2.5VL on each page OCR result, get vision description
                        qwen_result = None
                        vision_desc = None
                        if qwen_provider:
                            try:
                                # Use a vision prompt for description (not sensitivity)
                                vision_desc = qwen_provider.analyze_vision(img)
                                qwen_result = qwen_provider.analyze_text(ocr_result.strip())
                            except Exception as qe:
                                logger.error(f"Qwen2.5VL OCR analysis failed for {file_path} page {page_num}: {qe}")
                                qwen_result = {'error': str(qe)}
                                vision_desc = {'error': str(qe)}
                        qwen2_5vl_results.append({
                            'page': page_num,
                            'qwen2.5vl_result': qwen_result,
                            'vision_description': vision_desc
                        })
                        if vision_desc and isinstance(vision_desc, dict):
                            desc = vision_desc.get('explanation') or str(vision_desc)
                            vision_descriptions.append(f"[Vision description page {page_num}]: {desc}")
                    # Replace text_content with only vision descriptions for main model
                    if vision_descriptions:
                        text_content = "\n".join(vision_descriptions).strip()
                        extraction_method = (extraction_method or "pdf") + "+qwen2.5vl-vision"
                except Exception as e:
                    logger.warning(f"PDF OCR failed for {file_path}: {e}")
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass
            return {
                'text_content': text_content.strip(),
                'extraction_method': extraction_method or 'pdf',
                'confidence': 1.0 if text_content.strip() else 0.0,
                'metadata': {
                    'pdf_pages': pdf_pages,
                    'qwen2.5vl_results': qwen2_5vl_results if qwen2_5vl_results else None
                }
            }
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'pdf_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_word(self, file_path: Path) -> Dict:
        """
        Extract text from Word documents
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                'text_content': text_content.strip(),
                'extraction_method': 'python-docx',
                'confidence': 1.0,
                'metadata': {
                    'paragraphs': len(doc.paragraphs)
                }
            }
            
        except Exception as e:
            logger.error(f"Word document extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'word_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_text(self, file_path: Path) -> Dict:
        """
        Extract text from plain text files
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            return {
                'text_content': text_content.strip(),
                'extraction_method': 'native',
                'confidence': 1.0,
                'metadata': {
                    'encoding': 'utf-8'
                }
            }
            
        except Exception as e:
            logger.error(f"Text file extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'text_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_spreadsheet(self, file_path: Path) -> Dict:
        """
        Extract text from spreadsheet files
        
        Args:
            file_path: Path to spreadsheet file
            
        Returns:
            Dictionary with extracted text
        """
        file_extension = get_file_extension(str(file_path))
        
        if file_extension == '.csv':
            return self._extract_from_csv(file_path)
        elif file_extension in ['.xls', '.xlsx']:
            return self._extract_from_excel(file_path)
        else:
            return {
                'text_content': '',
                'extraction_method': 'unsupported',
                'error': f'Unsupported spreadsheet type: {file_extension}'
            }
    
    def _extract_from_csv(self, file_path: Path) -> Dict:
        """
        Extract text from CSV file
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            df = pd.read_csv(file_path)
            text_content = df.to_string(index=False)
            
            return {
                'text_content': text_content,
                'extraction_method': 'pandas',
                'confidence': 1.0,
                'metadata': {
                    'rows': len(df),
                    'columns': len(df.columns)
                }
            }
            
        except Exception as e:
            logger.error(f"CSV extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'csv_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_excel(self, file_path: Path) -> Dict:
        """
        Extract text from Excel file
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with extracted text
        """
        file_extension = get_file_extension(str(file_path))
        
        # Try openpyxl first (for .xlsx, .xlsm, etc.)
        if file_extension in ['.xlsx', '.xlsm', '.xltx', '.xltm']:
            try:
                workbook = load_workbook(file_path, data_only=True)
                text_content = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_content += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        if row_text.strip():
                            text_content += row_text + "\n"
                
                return {
                    'text_content': text_content.strip(),
                    'extraction_method': 'openpyxl',
                    'confidence': 1.0,
                    'metadata': {
                        'sheets': len(workbook.sheetnames),
                        'sheet_names': workbook.sheetnames
                    }
                }
                
            except Exception as e:
                logger.warning(f"openpyxl failed for {file_path}, trying pandas: {e}")
        
        # Try pandas for .xls files or as fallback
        try:
            # Try different engines for pandas
            engines = ['xlrd', 'openpyxl']
            text_content = ""
            metadata = {}
            
            for engine in engines:
                try:
                    # Read all sheets
                    excel_file = pd.ExcelFile(file_path, engine=engine)
                    metadata['sheets'] = len(excel_file.sheet_names)
                    metadata['sheet_names'] = excel_file.sheet_names
                    
                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, engine=engine)
                        text_content += f"\n--- Sheet: {sheet_name} ---\n"
                        text_content += df.to_string(index=False) + "\n"
                    
                    return {
                        'text_content': text_content.strip(),
                        'extraction_method': f'pandas_{engine}',
                        'confidence': 1.0,
                        'metadata': metadata
                    }
                    
                except Exception as e:
                    logger.debug(f"pandas with {engine} failed: {e}")
                    continue
            
            # If all engines failed, try basic text extraction
            logger.warning(f"All Excel engines failed for {file_path}, trying basic text extraction")
            return self._extract_excel_as_text(file_path)
            
        except Exception as e:
            logger.error(f"Excel extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'excel_failed',
                'confidence': 0.0,
                'error': str(e)
            }

    def _extract_excel_as_text(self, file_path: Path) -> Dict:
        """
        Fallback method to extract text from Excel files by treating them as ZIP archives
        (for .xlsx files) or trying to read as binary
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            text_content = ""
            
            # Try to read as ZIP (for .xlsx files)
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Look for shared strings and sheet data
                    if 'xl/sharedStrings.xml' in zip_file.namelist():
                        with zip_file.open('xl/sharedStrings.xml') as f:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            for si in root.findall('.//{*}t'):
                                if si.text:
                                    text_content += si.text + " "
                    
                    # Look for sheet data
                    sheet_files = [f for f in zip_file.namelist() if f.startswith('xl/worksheets/sheet') and f.endswith('.xml')]
                    for sheet_file in sheet_files[:3]:  # Limit to first 3 sheets
                        try:
                            with zip_file.open(sheet_file) as f:
                                tree = ET.parse(f)
                                root = tree.getroot()
                                for v in root.findall('.//{*}v'):
                                    if v.text:
                                        text_content += v.text + " "
                        except Exception:
                            continue
                            
            except zipfile.BadZipFile:
                # Not a ZIP file, try to read as binary and extract text
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                        # Try to decode as text (might work for some formats)
                        try:
                            text_content = content.decode('utf-8', errors='ignore')
                        except:
                            text_content = content.decode('latin-1', errors='ignore')
                except Exception:
                    pass
            
            if text_content.strip():
                return {
                    'text_content': text_content.strip(),
                    'extraction_method': 'excel_fallback',
                    'confidence': 0.5,
                    'metadata': {
                        'method': 'fallback_extraction'
                    }
                }
            else:
                return {
                    'text_content': '',
                    'extraction_method': 'excel_unsupported',
                    'confidence': 0.0,
                    'error': 'Excel file format not supported by any available method'
                }
                
        except Exception as e:
            logger.error(f"Excel fallback extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'excel_fallback_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def _extract_from_presentation(self, file_path: Path) -> Dict:
        """
        Extract text from presentation files (PowerPoint)
        
        Args:
            file_path: Path to presentation file
            
        Returns:
            Dictionary with extracted text
        """
        file_extension = get_file_extension(str(file_path))
        
        if file_extension in ['.ppt', '.pptx']:
            return self._extract_from_powerpoint(file_path)
        else:
            return {
                'text_content': '',
                'extraction_method': 'unsupported',
                'confidence': 0.0,
                'error': f'Unsupported presentation type: {file_extension}'
            }
    
    def _extract_from_powerpoint(self, file_path: Path) -> Dict:
        """
        Extract text from PowerPoint files
        
        Args:
            file_path: Path to PowerPoint file
            
        Returns:
            Dictionary with extracted text
        """
        try:
            # PowerPoint files are ZIP archives containing XML files
            # We'll extract text from the XML content
            import zipfile
            import xml.etree.ElementTree as ET
            
            text_content = ""
            slide_count = 0
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Look for slide files
                slide_files = [f for f in zip_file.namelist() if f.startswith('ppt/slides/slide') and f.endswith('.xml')]
                
                for slide_file in slide_files:
                    try:
                        with zip_file.open(slide_file) as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            
                            # Extract text from text elements
                            # PowerPoint uses different namespaces, so we'll search broadly
                            for elem in root.iter():
                                if elem.text and elem.text.strip():
                                    text_content += elem.text.strip() + "\n"
                            
                            slide_count += 1
                    except Exception as e:
                        logger.warning(f"Error processing slide {slide_file}: {e}")
                        continue
            
            if not text_content.strip():
                # Fallback: try to extract from content types and other XML files
                try:
                    with zipfile.ZipFile(file_path, 'r') as zip_file:
                        # Try to get text from any XML files
                        xml_files = [f for f in zip_file.namelist() if f.endswith('.xml')]
                        for xml_file in xml_files[:5]:  # Limit to first 5 files
                            try:
                                with zip_file.open(xml_file) as f:
                                    content = f.read().decode('utf-8', errors='ignore')
                                    # Simple text extraction from XML
                                    import re
                                    text_matches = re.findall(r'>([^<]+)<', content)
                                    for match in text_matches:
                                        if match.strip() and len(match.strip()) > 2:
                                            text_content += match.strip() + "\n"
                            except Exception:
                                continue
                except Exception as e:
                    logger.warning(f"Fallback PowerPoint extraction failed: {e}")
            
            return {
                'text_content': text_content.strip(),
                'extraction_method': 'powerpoint_xml',
                'confidence': 0.8 if text_content.strip() else 0.0,
                'metadata': {
                    'slides_processed': slide_count,
                    'file_type': 'powerpoint'
                }
            }
            
        except Exception as e:
            logger.error(f"PowerPoint extraction failed for {file_path}: {e}")
            return {
                'text_content': '',
                'extraction_method': 'powerpoint_failed',
                'confidence': 0.0,
                'error': str(e)
            }
    
    def extract_batch(self, file_paths: List[Path]) -> List[Dict]:
        """
        Extract text from multiple files
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of extraction results
        """
        results = []
        
        for file_path in file_paths:
            logger.info(f"Extracting text from {file_path}")
            result = self.extract_text(file_path)
            results.append(result)
        
        return results
    
    def get_extraction_summary(self, results: List[Dict]) -> Dict:
        """
        Get summary of text extraction results
        
        Args:
            results: List of extraction results
            
        Returns:
            Summary dictionary
        """
        total = len(results)
        successful = sum(1 for r in results if r.get('text_content', '').strip())
        failed = sum(1 for r in results if r.get('error'))
        
        methods = {}
        for result in results:
            method = result.get('extraction_method', 'unknown')
            methods[method] = methods.get(method, 0) + 1
        
        return {
            'total_files': total,
            'successful_extractions': successful,
            'failed_extractions': failed,
            'success_rate': successful / total if total > 0 else 0,
            'extraction_methods': methods
        } 